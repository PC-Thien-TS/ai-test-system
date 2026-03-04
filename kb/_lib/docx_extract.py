from __future__ import annotations

import re
from pathlib import Path

from kb._lib.types import DocumentBlock


def extract_document_blocks(path: Path) -> list[DocumentBlock]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".md":
        return _extract_markdown(path)
    if suffix == ".txt":
        return _extract_text(path)
    raise ValueError(f"Unsupported document type: {path}")


def _extract_docx(path: Path) -> list[DocumentBlock]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for DOCX extraction. Install it locally with 'pip install python-docx'."
        ) from exc

    document = Document(path)
    headings: list[str] = []
    blocks: list[DocumentBlock] = []
    block_index = 0

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        heading_level = _heading_level(style_name)
        if heading_level is not None:
            while len(headings) >= heading_level:
                headings.pop()
            headings.append(text)
            continue
        blocks.append(
            DocumentBlock(
                source_file=path.name,
                text=text,
                headings=list(headings),
                block_index=block_index,
            )
        )
        block_index += 1

    for table in document.tables:
        table_lines: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_lines.append(" | ".join(cells))
        if table_lines:
            blocks.append(
                DocumentBlock(
                    source_file=path.name,
                    text="\n".join(table_lines),
                    headings=list(headings),
                    block_index=block_index,
                )
            )
            block_index += 1
    return blocks


def _extract_markdown(path: Path) -> list[DocumentBlock]:
    headings: list[str] = []
    blocks: list[DocumentBlock] = []
    buffer: list[str] = []
    block_index = 0

    def flush() -> None:
        nonlocal block_index
        text = "\n".join(buffer).strip()
        if not text:
            return
        blocks.append(
            DocumentBlock(
                source_file=path.name,
                text=text,
                headings=list(headings),
                block_index=block_index,
            )
        )
        buffer.clear()
        block_index += 1

    for line in path.read_text(encoding="utf-8", errors="replace").splitlines():
        match = re.match(r"^(#{1,6})\s+(.*)$", line.strip())
        if match:
            flush()
            level = len(match.group(1))
            title = match.group(2).strip()
            while len(headings) >= level:
                headings.pop()
            headings.append(title)
            continue
        if not line.strip():
            flush()
            continue
        buffer.append(line.rstrip())
    flush()
    return blocks


def _extract_text(path: Path) -> list[DocumentBlock]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    parts = [part.strip() for part in re.split(r"\n\s*\n", raw) if part.strip()]
    return [
        DocumentBlock(
            source_file=path.name,
            text=part,
            headings=[],
            block_index=index,
        )
        for index, part in enumerate(parts)
    ]


def _heading_level(style_name: str) -> int | None:
    match = re.search(r"heading\s+([1-6])", style_name or "", flags=re.IGNORECASE)
    if not match:
        return None
    return int(match.group(1))
